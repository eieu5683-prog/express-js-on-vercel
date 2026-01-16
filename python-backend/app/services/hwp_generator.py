"""
HWP 파일 생성 서비스
pyhwp 또는 python-docx를 사용하여 PSST 데이터를 HWP 양식에 삽입
"""
import os
import tempfile
from pathlib import Path
from typing import Optional
from app.models.psst import PSSTDocument


class HWPGenerator:
    """HWP 파일 생성기"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        초기화
        
        Args:
            template_path: HWP 템플릿 파일 경로 (선택사항)
        """
        self.template_path = template_path or self._get_default_template_path()
        self.output_dir = Path(tempfile.gettempdir()) / "psst_hwp"
        self.output_dir.mkdir(exist_ok=True)
    
    def _get_default_template_path(self) -> Optional[str]:
        """기본 템플릿 경로 반환"""
        # 템플릿 파일이 있으면 사용, 없으면 None
        template_file = Path(__file__).parent.parent.parent / "templates" / "psst_template.hwp"
        if template_file.exists():
            return str(template_file)
        return None
    
    async def generate_hwp(self, psst_doc: PSSTDocument) -> str:
        """
        PSST 문서를 HWP 파일로 생성
        
        Args:
            psst_doc: PSST 문서 객체
            
        Returns:
            생성된 HWP 파일 경로
        """
        try:
            # pyhwp를 사용한 HWP 생성 시도
            return await self._generate_with_pyhwp(psst_doc)
        except ImportError:
            # pyhwp가 없으면 python-docx로 DOCX 생성 후 안내
            return await self._generate_with_docx(psst_doc)
    
    async def _generate_with_pyhwp(self, psst_doc: PSSTDocument) -> str:
        """
        pyhwp를 사용하여 HWP 파일 생성
        
        Args:
            psst_doc: PSST 문서 객체
            
        Returns:
            생성된 HWP 파일 경로
        """
        try:
            from pyhwp import hwp5
            from pyhwp.hwp5 import plat
            
            # 출력 파일 경로
            output_path = self.output_dir / f"psst_{psst_doc.metadata.get('industryName', 'document')}.hwp"
            
            # 템플릿이 있으면 복사, 없으면 새로 생성
            if self.template_path and os.path.exists(self.template_path):
                import shutil
                shutil.copy(self.template_path, output_path)
                hwp_file = hwp5.HWP5File(output_path)
            else:
                # 새 HWP 파일 생성
                hwp_file = hwp5.HWP5File(output_path, mode='w')
            
            # HWP 파일에 PSST 데이터 삽입
            self._insert_psst_data_to_hwp(hwp_file, psst_doc)
            
            hwp_file.close()
            return str(output_path)
            
        except ImportError:
            raise ImportError("pyhwp 라이브러리가 설치되지 않았습니다. pip install pyhwp를 실행하세요.")
        except Exception as e:
            raise Exception(f"HWP 생성 중 오류 발생: {str(e)}")
    
    async def _generate_with_docx(self, psst_doc: PSSTDocument) -> str:
        """
        python-docx를 사용하여 DOCX 파일 생성 (HWP 대체)
        
        Args:
            psst_doc: PSST 문서 객체
            
        Returns:
            생성된 DOCX 파일 경로
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 출력 파일 경로
            output_path = self.output_dir / f"psst_{psst_doc.metadata.get('industryName', 'document')}.docx"
            
            # 템플릿이 있으면 열기, 없으면 새로 생성
            if self.template_path and os.path.exists(self.template_path):
                doc = Document(self.template_path)
            else:
                doc = Document()
            
            # PSST 데이터 삽입
            self._insert_psst_data_to_docx(doc, psst_doc)
            
            # 파일 저장
            doc.save(str(output_path))
            return str(output_path)
            
        except ImportError:
            raise ImportError("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx를 실행하세요.")
        except Exception as e:
            raise Exception(f"DOCX 생성 중 오류 발생: {str(e)}")
    
    def _insert_psst_data_to_hwp(self, hwp_file, psst_doc: PSSTDocument):
        """
        HWP 파일에 PSST 데이터 삽입
        
        Args:
            hwp_file: pyhwp HWP5File 객체
            psst_doc: PSST 문서 객체
        """
        try:
            from pyhwp.hwp5 import plat
            
            # HWP 파일 구조에 맞춰 데이터 삽입
            # 템플릿이 있는 경우 특정 필드에 매핑, 없으면 새로 생성
            
            problem = psst_doc.problem
            solution = psst_doc.solution
            scale_up = psst_doc.scale_up
            team = psst_doc.team
            metadata = psst_doc.metadata
            
            # 제목 설정
            # hwp_file.body.section[0].paragraphs[0].text = "PSST 사업계획서"
            
            # 메타데이터 삽입
            # 업종 정보: metadata['industryName'], metadata['industryCode']
            
            # Problem 섹션 삽입
            # 시장 문제점: problem['marketIssues']
            # 사회적 이유: problem['socialReasons']
            # 경제적 이유: problem['economicReasons']
            # 시급성: problem['urgency']
            
            # Solution 섹션 삽입
            # 핵심 기술: solution['coreTechnology']
            # 주요 기능: solution['keyFeatures']
            # 차별화 포인트: solution['differentiation']
            # 경쟁 우위: solution['competitiveAdvantage']
            
            # Scale-up 섹션 삽입
            # 수익 모델: scale_up['revenueModel']
            # 수익원: scale_up['revenueStreams']
            # 시장 진입 전략: scale_up['marketEntryStrategy']
            # 확장 계획: scale_up['expansionPlan']
            # 시장 점유율 목표: scale_up['marketShareGoal']
            # 마일스톤: scale_up['milestones']
            
            # Team 섹션 삽입
            # 대표자: team['ceo']
            # 핵심 팀원: team['coreTeam']
            # 네트워크: team['network']
            # 팀 역량: team['capabilities']
            
            # AI 표기법 준수 문구 삽입 (문서 하단)
            # "본 초안은 나랏돈네비 AI 기술을 활용하여 작성되었습니다."
            # "(2026년 1월 23일부터 시행되는 AI 생성 콘텐츠 표기 의무화 규정 준수)"
            
            # 주의: pyhwp의 실제 API는 버전에 따라 다를 수 있으므로
            # 실제 사용 시 pyhwp 문서를 참고하여 구현해야 합니다.
            
            pass
        except Exception as e:
            raise Exception(f"HWP 데이터 삽입 중 오류: {str(e)}")
    
    def _insert_psst_data_to_docx(self, doc, psst_doc: PSSTDocument):
        """
        DOCX 파일에 PSST 데이터 삽입
        
        Args:
            doc: python-docx Document 객체
            psst_doc: PSST 문서 객체
        """
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 제목
        title = doc.add_heading('PSST 사업계획서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 메타데이터
        metadata = psst_doc.metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"업종: {metadata.get('industryName', '')} ({metadata.get('industryCode', '')})").bold = True
        doc.add_paragraph(f"생성일: {metadata.get('createdAt', '')}")
        doc.add_paragraph()  # 빈 줄
        
        # 1. Problem
        doc.add_heading('1. Problem (문제 인식)', level=1)
        problem = psst_doc.problem
        
        doc.add_heading('시장의 문제점', level=2)
        for issue in problem.get('marketIssues', []):
            doc.add_paragraph(issue, style='List Bullet')
        
        doc.add_heading('사회적 이유', level=2)
        for reason in problem.get('socialReasons', []):
            doc.add_paragraph(reason, style='List Bullet')
        
        doc.add_heading('경제적 이유', level=2)
        for reason in problem.get('economicReasons', []):
            doc.add_paragraph(reason, style='List Bullet')
        
        doc.add_heading('해결의 시급성', level=2)
        doc.add_paragraph(problem.get('urgency', ''))
        doc.add_paragraph()  # 빈 줄
        
        # 2. Solution
        doc.add_heading('2. Solution (해결 방안)', level=1)
        solution = psst_doc.solution
        
        doc.add_heading('핵심 기술', level=2)
        doc.add_paragraph(solution.get('coreTechnology', ''))
        
        doc.add_heading('주요 기능', level=2)
        for feature in solution.get('keyFeatures', []):
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_heading('경쟁사 대비 차별화 포인트', level=2)
        for point in solution.get('differentiation', []):
            doc.add_paragraph(point, style='List Bullet')
        
        doc.add_heading('경쟁 우위', level=2)
        doc.add_paragraph(solution.get('competitiveAdvantage', ''))
        doc.add_paragraph()  # 빈 줄
        
        # 3. Scale-up
        doc.add_heading('3. Scale-up (성장 전략)', level=1)
        scale_up = psst_doc.scale_up
        
        doc.add_heading('수익 창출 방안', level=2)
        doc.add_paragraph(scale_up.get('revenueModel', ''))
        
        doc.add_heading('수익원', level=2)
        for stream in scale_up.get('revenueStreams', []):
            doc.add_paragraph(stream, style='List Bullet')
        
        doc.add_heading('시장 진입 전략', level=2)
        doc.add_paragraph(scale_up.get('marketEntryStrategy', ''))
        
        doc.add_heading('확장 계획', level=2)
        doc.add_paragraph(scale_up.get('expansionPlan', ''))
        
        doc.add_heading('3년 내 시장 점유율 목표', level=2)
        doc.add_paragraph(scale_up.get('marketShareGoal', ''))
        
        doc.add_heading('주요 마일스톤', level=2)
        for milestone in scale_up.get('milestones', []):
            milestone_text = f"{milestone.get('year', '')}년 {milestone.get('quarter', '')}분기: {milestone.get('goal', '')} ({milestone.get('metric', '')})"
            doc.add_paragraph(milestone_text, style='List Bullet')
        doc.add_paragraph()  # 빈 줄
        
        # 4. Team
        doc.add_heading('4. Team (팀 구성)', level=1)
        team = psst_doc.team
        
        doc.add_heading('대표자 (CEO)', level=2)
        ceo = team.get('ceo', {})
        doc.add_paragraph(f"이름: {ceo.get('name', '')}")
        doc.add_paragraph(f"역할: {ceo.get('role', '')}")
        doc.add_paragraph(f"전문 분야: {', '.join(ceo.get('expertise', []))}")
        doc.add_paragraph(f"경력: {ceo.get('experience', '')}")
        if ceo.get('education'):
            doc.add_paragraph(f"학력: {ceo.get('education', '')}")
        
        doc.add_heading('핵심 팀원', level=2)
        for member in team.get('coreTeam', []):
            doc.add_paragraph(f"{member.get('name', '')} ({member.get('role', '')})", style='Heading 3')
            doc.add_paragraph(f"전문 분야: {', '.join(member.get('expertise', []))}")
            doc.add_paragraph(f"경력: {member.get('experience', '')}")
            if member.get('education'):
                doc.add_paragraph(f"학력: {member.get('education', '')}")
        
        doc.add_heading('네트워크 및 파트너십', level=2)
        for network in team.get('network', []):
            doc.add_paragraph(network, style='List Bullet')
        
        doc.add_heading('팀 역량', level=2)
        for capability in team.get('capabilities', []):
            doc.add_paragraph(capability, style='List Bullet')
        
        # AI 표기법 준수 문구
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run('본 초안은 나랏돈네비 AI 기술을 활용하여 작성되었습니다.')
        run.bold = True
        run.font.size = Pt(10)
        
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run = disclaimer_para.add_run('(2026년 1월 23일부터 시행되는 AI 생성 콘텐츠 표기 의무화 규정 준수)')
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)

